"""
A module that extracts information from files
"""
import logging
import io
import boto3
import PyPDF2
import docx

logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)
logger.setLevel(level=logging.WARNING)

s3 = boto3.client('s3')

def get_pdf_text_from_path(filepath : str) -> str:
    """Extract text in a pdf file
    
    Arguments:
        filepath {str} -- path to the pdf file
    
    Returns:
        str -- text in the pdf file
    """
    with open(filepath, 'rb') as file_obj:
        reader = PyPDF2.PdfFileReader(file_obj)
        text = []
        for i in range(reader.numPages):
            page_obj = reader.getPage(i)
            text.append(page_obj.extractText())
        return "\n".join(text).strip()

def get_pdf_text_from_binary_data(binary_data : io.BytesIO) -> str:
    """Extract text in a pdf file given byte streams
    
    Arguments:
        binary_data {io.BytesIO} -- pdf content in bytes
    
    Returns:
        str -- text in the pdf file
    """
    try:
        reader = PyPDF2.PdfFileReader(binary_data)
        text = []
        for i in range(reader.numPages):
            page_obj = reader.getPage(i)
            text.append(page_obj.extractText())
        return "\n".join(text).strip()
    except Exception as e:
        logger.error(e)
        raise e

def get_docx_text_from_path(filepath : str) -> str:
    """Extract text in a docx file
    
    Arguments:
        filepath {str} -- path to the docx file
    
    Returns:
        str -- text in the docx file
    """
    doc = docx.Document(filepath)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return "\n".join(text).strip()

def get_docx_text_from_binary_data(binary_data : io.BytesIO) -> str:
    """Extract text in a docx file from binary data
    
    Arguments:
        binary_data {io.BytesIO} -- docx content in bytes
    
    Returns:
        str -- text in the docx file
    """
    try:
        doc = docx.Document(binary_data)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(text).strip()
    except Exception as e:
        logger.error(e)
        raise e

def get_txt_text_from_path(filepath : str) -> str:
    """Extract text in a txt file
    
    Arguments:
        filepath {str} -- path to the txt file
    
    Returns:
        str -- text in the txt file
    """
    with open(filepath, 'r') as f:
        text = f.read()
    return text.strip()

def get_txt_text_from_binary_data(binary_data : io.BytesIO) -> str:
    """Extract text in a txt file from binary data
    
    Arguments:
        binary_data {io.BytesIO} -- txt content in bytes
    
    Returns:
        str -- text in the txt file
    """
    try:
        return binary_data.read().decode("UTF-8").strip()
    except Exception as e:
        logger.error(e)
        raise e

def get_file_text_from_path(filepath : str) -> str:
    """Extract the text from a text file
    
    Arguments:
        filepath {str} -- path to the text file
    
    Returns:
        str -- text in the text file
    """
    extension = filepath.split('.')[-1]
    if extension == 'pdf':
        return get_pdf_text_from_path(filepath)
    elif extension == 'docx':
        return get_docx_text_from_path(filepath)
    elif extension == 'txt':
        return get_txt_text_from_path(filepath)

def get_file_text_from_binary_data(extension : str, binary_data : io.BytesIO) -> str:
    """Extract the text from a text file given its binary data
    
    Arguments:
        extension {str} -- file extension
        binary_data {io.BytesIO} -- text content in binary format
    
    Returns:
        str -- text in the text file
    """
    if extension == 'pdf':
        return get_pdf_text_from_binary_data(binary_data)
    elif extension == 'docx':
        return get_docx_text_from_binary_data(binary_data)
    elif extension == 'txt':
        return get_txt_text_from_binary_data(binary_data)

def get_binary_data_from_file_in_s3(bucket : str, key : str) -> io.BytesIO:
    """Extract binary data from file in an S3 bucket
    
    Arguments:
        bucket {str} -- S3 Bucket Name
        key {str} -- File Key
    
    Returns:
        io.BytesIO -- binary data of the file
    """

    try:
        response = s3.get_object(Bucket=bucket, Key=key)
        s3_file_content = response['Body'].read()
        return io.BytesIO(s3_file_content)
    except Exception as e:
        logger.error(f"Unable to read from bucket and {bucket} key {key}")
        logger.error(e)
        raise e
